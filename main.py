# Import required libraries from FastAPI and others
# Import required libraries from FastAPI and others
import random
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from fastapi.responses import StreamingResponse
from fastapi import FastAPI, Request, Form, HTTPException, status, Depends
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.middleware.sessions import SessionMiddleware
import requests
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches
from datetime import datetime
from reportlab.lib.units import inch


import csv
from bs4 import BeautifulSoup

import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors

import asyncpg
import os
from dotenv import load_dotenv

load_dotenv()

app = FastAPI()
app.add_middleware(
    SessionMiddleware,
    secret_key="tut123",
    max_age=3600  # 1 hour
)

app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

import ssl

DATABASE_URL = os.getenv(
    "DATABASE_URL",
    "postgresql://edubridge_db_58vd_user:XIScMCDoEJHxq5JJmvA0iAAiw91NHy3C@dpg-d18jk8mmcj7s73a26510-a.oregon-postgres.render.com:5432/edubridge_db_58vd?sslmode=require"
)



db_pool = None

@app.on_event("startup")
async def startup():
    global db_pool
    db_pool = await asyncpg.create_pool(DATABASE_URL, timeout=30)

@app.on_event("shutdown")
async def shutdown():
    if db_pool:
        await db_pool.close()

# ------------------ Page Routes ------------------
@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/about", response_class=HTMLResponse)
async def about_page(request: Request):
    return templates.TemplateResponse("about.html", {"request": request})

@app.get("/services", response_class=HTMLResponse)
async def services_page(request: Request):
    return templates.TemplateResponse("services.html", {"request": request})

@app.get("/contact", response_class=HTMLResponse)
async def contact_form(request: Request, success: int = 0):
    return templates.TemplateResponse("contact.html", {"request": request, "success": bool(success)})

@app.post("/contact", response_class=HTMLResponse)
async def contact_submit(request: Request, name: str = Form(...), email: str = Form(...), message: str = Form(...)):
    print(f"Contact Submission:\nName: {name}\nEmail: {email}\nMessage: {message}")
    return RedirectResponse(url="/contact?success=1", status_code=status.HTTP_303_SEE_OTHER)

# ------------------- UJ LOGIN -------------------
@app.get("/uj-login", response_class=HTMLResponse)
async def uj_login(request: Request):
    return templates.TemplateResponse("uj-login.html", {"request": request})

@app.post("/uj-login", response_class=HTMLResponse)
async def uj_login_post(request: Request, studentID: str = Form(...), password: str = Form(...)):
    async with db_pool.acquire() as conn:
        uj_id = await conn.fetchval("""
            SELECT university_id FROM university 
            WHERE name ILIKE 'University of Johannesburg'
        """)
        student = await conn.fetchrow("""
            SELECT * FROM student 
            WHERE student_id = $1 AND password = $2 AND home_university = $3
        """, int(studentID), password, uj_id)
        if not student:
            return templates.TemplateResponse("uj-login.html", {
                "request": request,
                "error": "Invalid UJ student ID or password."
            })
        request.session["student_id"] = studentID
        request.session["password"] = password
        request.session["home_university"] = "University of Johannesburg"
    return RedirectResponse(url="/tut-dashboard", status_code=status.HTTP_303_SEE_OTHER)

@app.get("/register-uj", response_class=HTMLResponse)
async def register_uj_get(request: Request):
    return templates.TemplateResponse("register-uj.html", {"request": request})

@app.post("/register-uj", response_class=HTMLResponse)
async def register_uj_student(
    request: Request,
    studentID: str = Form(...),
    name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    gender: str = Form(...)
):
    # Validate studentID
    if not studentID.isdigit() or len(studentID) != 11:
        return templates.TemplateResponse("register-uj.html", {
            "request": request,
            "error": "Student ID must be exactly 11 digits."
        })

    # Validate gender
    if gender not in ["Male", "Female"]:
        return templates.TemplateResponse("register-uj.html", {
            "request": request,
            "error": "Gender must be either Male or Female."
        })

    async with db_pool.acquire() as conn:
        # Check for existing student
        existing = await conn.fetchval("SELECT COUNT(*) FROM student WHERE student_id = $1", int(studentID))
        if existing:
            return templates.TemplateResponse("register-uj.html", {
                "request": request,
                "error": "Student already registered."
            })

        # Get or insert UJ university ID
        uj_id = await conn.fetchval("SELECT university_id FROM university WHERE name ILIKE 'University of Johannesburg'")
        if not uj_id:
            uj_id = await conn.fetchval("INSERT INTO university (name) VALUES ('University of Johannesburg') RETURNING university_id")

        # Insert new student
        await conn.execute("""
            INSERT INTO student (student_id, name, email, password, home_university, gender)
            VALUES ($1, $2, $3, $4, $5, $6)
        """, int(studentID), name, email, password, uj_id, gender)

    return RedirectResponse(url="/uj-login", status_code=status.HTTP_303_SEE_OTHER)

# ------------------- UP LOGIN -------------------
@app.get("/up-login", response_class=HTMLResponse)
async def up_login(request: Request):
    return templates.TemplateResponse("up_login.html", {"request": request})

@app.post("/up-login", response_class=HTMLResponse)
async def up_login_post(
    request: Request,
    studentID: str = Form(...),
    password: str = Form(...)
):
    async with db_pool.acquire() as conn:
        up_id = await conn.fetchval("""
            SELECT university_id FROM university 
            WHERE name ILIKE 'University of Pretoria'
        """)
        student = await conn.fetchrow("""
            SELECT * FROM student 
            WHERE student_id = $1 AND password = $2 AND home_university = $3
        """, int(studentID), password, up_id)
        if not student:
            return templates.TemplateResponse("up_login.html", {
                "request": request,
                "error": "Invalid UP student ID or password."
            })
        request.session["student_id"] = studentID
        request.session["password"] = password
        request.session["home_university"] = "University of Pretoria"
    return RedirectResponse(url="/tut-dashboard", status_code=status.HTTP_303_SEE_OTHER)

@app.get("/register-up", response_class=HTMLResponse)
async def register_up_get(request: Request):
    return templates.TemplateResponse("register-up.html", {"request": request})

@app.post("/register-up", response_class=HTMLResponse)
async def register_up_student(
    request: Request,
    studentID: str = Form(...),
    name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    gender: str = Form(...)
):
    # Validate studentID
    if not studentID.isdigit() or len(studentID) != 11:
        return templates.TemplateResponse("register-up.html", {
            "request": request,
            "error": "Student ID must be exactly 11 digits."
        })

    # Validate gender
    if gender not in ["Male", "Female"]:
        return templates.TemplateResponse("register-up.html", {
            "request": request,
            "error": "Gender must be either Male or Female."
        })

    async with db_pool.acquire() as conn:
        # Check for existing student
        existing = await conn.fetchval("SELECT COUNT(*) FROM student WHERE student_id = $1", int(studentID))
        if existing:
            return templates.TemplateResponse("register-up.html", {
                "request": request,
                "error": "Student already registered."
            })

        # Get or insert UP university ID
        up_id = await conn.fetchval("SELECT university_id FROM university WHERE name ILIKE 'University of Pretoria'")
        if not up_id:
            up_id = await conn.fetchval("INSERT INTO university (name) VALUES ('University of Pretoria') RETURNING university_id")

        # Insert new student
        await conn.execute("""
            INSERT INTO student (student_id, name, email, password, home_university, gender)
            VALUES ($1, $2, $3, $4, $5, $6)
        """, int(studentID), name, email, password, up_id, gender)

    return RedirectResponse(url="/up-login", status_code=status.HTTP_303_SEE_OTHER)

# ------------------- WITS LOGIN -------------------
@app.get("/wits-login", response_class=HTMLResponse)
async def wits_login(request: Request):
    return templates.TemplateResponse("wits_login.html", {"request": request})

@app.post("/wits-login", response_class=HTMLResponse)
async def wits_login_post(
    request: Request,
    studentID: str = Form(...),
    password: str = Form(...)
):
    async with db_pool.acquire() as conn:
        wits_id = await conn.fetchval("""
            SELECT university_id FROM university 
            WHERE name ILIKE 'University of the Witwatersrand'
        """)
        student = await conn.fetchrow("""
            SELECT * FROM student 
            WHERE student_id = $1 AND password = $2 AND home_university = $3
        """, int(studentID), password, wits_id)
        if not student:
            return templates.TemplateResponse("wits_login.html", {
                "request": request,
                "error": "Invalid WITS student ID or password."
            })
        request.session["student_id"] = studentID
        request.session["password"] = password
        request.session["home_university"] = "University of the Witwatersrand"
    return RedirectResponse(url="/tut-dashboard", status_code=status.HTTP_303_SEE_OTHER)

@app.get("/register-wits", response_class=HTMLResponse)
async def register_wits_get(request: Request):
    return templates.TemplateResponse("register-wits.html", {"request": request})

@app.post("/register-wits", response_class=HTMLResponse)
async def register_wits_student(
    request: Request,
    studentID: str = Form(...),
    name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    gender: str = Form(...)
):
    # Validate studentID
    if not studentID.isdigit() or len(studentID) != 11:
        return templates.TemplateResponse("register-wits.html", {
            "request": request,
            "error": "Student ID must be exactly 11 digits."
        })

    # Validate gender
    if gender not in ["Male", "Female"]:
        return templates.TemplateResponse("register-wits.html", {
            "request": request,
            "error": "Gender must be either Male or Female."
        })

    async with db_pool.acquire() as conn:
        # Check for existing student
        existing = await conn.fetchval("SELECT COUNT(*) FROM student WHERE student_id = $1", int(studentID))
        if existing:
            return templates.TemplateResponse("register-wits.html", {
                "request": request,
                "error": "Student already registered."
            })

        # Get or insert WITS university ID
        wits_id = await conn.fetchval("SELECT university_id FROM university WHERE name ILIKE 'University of the Witwatersrand'")
        if not wits_id:
            wits_id = await conn.fetchval("INSERT INTO university (name) VALUES ('University of the Witwatersrand') RETURNING university_id")

        # Insert new student
        await conn.execute("""
            INSERT INTO student (student_id, name, email, password, home_university, gender)
            VALUES ($1, $2, $3, $4, $5, $6)
        """, int(studentID), name, email, password, wits_id, gender)

    return RedirectResponse(url="/wits-login", status_code=status.HTTP_303_SEE_OTHER)

# ------------------- UNIVERSITY SELECTION -------------------
@app.get("/university-selection", response_class=HTMLResponse)
async def university_selection(request: Request):
    return templates.TemplateResponse("university-selection.html", {"request": request})

@app.post("/university-selection")
async def university_selection_post(request: Request, universitySelect: str = Form(...)):
    if universitySelect == "Tshwane University of Technology (TUT)":
        return RedirectResponse(url="/tut-login?universitySelect=TUT", status_code=302)
    elif universitySelect == "University of Johannesburg":
        return RedirectResponse(url="/uj-login?universitySelect=UJ", status_code=302)
    elif universitySelect == "University of Pretoria (UP)":
        return RedirectResponse(url="/up-login?universitySelect=UP", status_code=302)
    elif universitySelect == "University of the Witwatersrand (WITS)":
        return RedirectResponse(url="/wits-login?universitySelect=WITS", status_code=302)
    else:
        return templates.TemplateResponse("university-selection.html", {
            "request": request,
            "error": "Selected university is not yet supported."
        })

# ------------------- MODULE SELECTION -------------------
@app.get("/modules", response_class=HTMLResponse)
async def modules_page(request: Request):
    student_id = request.session.get("student_id")
    if not student_id:
        return RedirectResponse(url="/university-selection", status_code=302)
    
    async with db_pool.acquire() as conn:
        # Get student details
        student = await conn.fetchrow("""
            SELECT s.*, u.name AS university_name
            FROM student s
            JOIN university u ON s.home_university = u.university_id
            WHERE s.student_id = $1
        """, int(student_id))
        
        # Get registered modules
        registered_modules = await conn.fetch("""
            SELECT m.name FROM registration r
            JOIN module m ON r.module_id = m.module_id
            WHERE r.student_id = $1
        """, int(student_id))
    
    registered_module_names = {mod["name"] for mod in registered_modules}
    
    # Rest of your logic (university-specific modules)
    university = student["university_name"].lower()
    
    if "tshwane" in university:
        modules = [
            {"name": "Artificial Intelligence", "status": "register"},
            {"name": "Internet Programming", "status": "offered"},
            {"name": "Mobile Programming", "status": "offered"},
            {"name": "Database Programming", "status": "offered"},
            {"name": "Web Server Management", "status": "register"},
            {"name": "Distributed System", "status": "register"},
            {"name": "Software Project", "status": "offered"},
        ]
    elif "johannesburg" in university:
        modules = [
            {"name": "Artificial Intelligence", "status": "offered"},
            {"name": "Internet Programming", "status": "register"},
            {"name": "Mobile Programming", "status": "register"},
            {"name": "Database Programming", "status": "register"},
            {"name": "Web Server Management", "status": "offered"},
            {"name": "Distributed System", "status": "offered"},
            {"name": "Software Project", "status": "register"},
        ]
    elif "pretoria" in university:
        modules = [
            {"name": "Artificial Intelligence", "status": "offered"},
            {"name": "Internet Programming", "status": "register"},
            {"name": "Mobile Programming", "status": "register"},
            {"name": "Database Programming", "status": "register"},
            {"name": "Web Server Management", "status": "offered"},
            {"name": "Distributed System", "status": "offered"},
            {"name": "Software Project", "status": "register"},
        ]
    elif "witwatersrand" in university:
        modules = [
            {"name": "Artificial Intelligence", "status": "offered"},
            {"name": "Internet Programming", "status": "register"},
            {"name": "Mobile Programming", "status": "register"},
            {"name": "Database Programming", "status": "register"},
            {"name": "Web Server Management", "status": "offered"},
            {"name": "Distributed System", "status": "offered"},
            {"name": "Software Project", "status": "register"},
        ]
    else:
        modules = [
            {"name": "Artificial Intelligence", "status": "register"},
            {"name": "Internet Programming", "status": "offered"},
            {"name": "Mobile Programming", "status": "offered"},
            {"name": "Database Programming", "status": "offered"},
            {"name": "Web Server Management", "status": "register"},
            {"name": "Distributed System", "status": "register"},
            {"name": "Software Project", "status": "offered"},
        ]
    
    # Update module status if already registered
    for module in modules:
        if module["name"] in registered_module_names:
            module["status"] = "registered"  # New status
    
    return templates.TemplateResponse("modules.html", {
        "request": request,
        "modules": modules,
        "student_university_name": student["university_name"],
        "message": "Register for modules not offered at your university.",
    })
# ------------------- TUT LOGIN -------------------
@app.get("/tut-login", response_class=HTMLResponse)
async def tut_login(request: Request):
    return templates.TemplateResponse("tut-login.html", {"request": request})

@app.post("/tut-login", response_class=HTMLResponse)
async def tut_login_post(request: Request, studentID: str = Form(...), password: str = Form(...)):
    async with db_pool.acquire() as conn:
        student = await conn.fetchrow("""
            SELECT * FROM student s
            JOIN university u ON s.home_university = u.university_id
            WHERE s.student_id = $1 AND s.password = $2 AND u.name ILIKE 'TUT'
        """, int(studentID), password)
        if not student:
            return templates.TemplateResponse("tut-login.html", {
                "request": request,
                "error": "Invalid student ID or password."
            })
        request.session["student_id"] = studentID
        request.session["password"] = password
    return RedirectResponse(url="/tut-dashboard", status_code=303)

@app.get("/tut-dashboard", response_class=HTMLResponse)
async def tut_dashboard(request: Request):
    return templates.TemplateResponse("tut-dashboard.html", {"request": request})

@app.get("/faculties", response_class=HTMLResponse)
async def faculties(request: Request):
    return templates.TemplateResponse("faculties.html", {"request": request})

@app.post("/select-faculty", response_class=HTMLResponse)
async def select_faculty(request: Request, facultyName: str = Form(...)):
    async with db_pool.acquire() as conn:
        record = await conn.fetchrow("SELECT faculty_id FROM faculty WHERE name = $1", facultyName)
        if not record:
            await conn.execute("INSERT INTO faculty (name, university_id) VALUES ($1, 1)", facultyName)
    return RedirectResponse(url="/courses", status_code=status.HTTP_303_SEE_OTHER)

@app.get("/courses", response_class=HTMLResponse)
async def courses(request: Request):
    return templates.TemplateResponse("courses.html", {"request": request})

# ------------------- ADMIN ROUTES -------------------
@app.get("/admin", response_class=HTMLResponse)
async def admin_dashboard(request: Request):
    await verify_admin(request)
    return templates.TemplateResponse("admin_dashboard.html", {"request": request})

@app.get("/admin/login", response_class=HTMLResponse)
async def admin_login(request: Request):
    return templates.TemplateResponse("admin_login.html", {"request": request})

@app.post("/admin/login", response_class=HTMLResponse)
async def admin_login_post(request: Request, adminEmail: str = Form(...), adminPassword: str = Form(...)):
    async with db_pool.acquire() as conn:
        admin = await conn.fetchrow("SELECT * FROM admin WHERE email = $1 AND password = $2", adminEmail, adminPassword)
        if not admin:
            return templates.TemplateResponse("admin_login.html", {"request": request, "error": "Invalid email or password"})
        request.session["admin_logged_in"] = True
    return RedirectResponse(url="/admin/dashboard", status_code=status.HTTP_303_SEE_OTHER)

@app.get("/admin/logout", response_class=HTMLResponse)
async def admin_logout(request: Request):
    request.session.clear()
    return RedirectResponse(url="/admin/login", status_code=status.HTTP_303_SEE_OTHER)

@app.get("/admin/dashboard", response_class=HTMLResponse)
async def admin_dashboard(request: Request):
    if not request.session.get("admin_logged_in"):
        return RedirectResponse(url="/admin/login", status_code=302)
    return templates.TemplateResponse("admin_dashboard.html", {"request": request})

@app.get("/admin/students", response_class=HTMLResponse)
async def view_students(
    request: Request,
    query: str = "",
    gender: str = "",
    university: str = "",
    student_id: str = ""
):
    await verify_admin(request)

    async with db_pool.acquire() as conn:
        # Fetch all universities for the dropdown
        universities = await conn.fetch("SELECT name FROM university ORDER BY name")

    # Fetch filtered students
    students = await fetch_filtered_students(query, gender, university, student_id)

    return templates.TemplateResponse("manage_students.html", {
        "request": request,
        "students": students,
        "universities": universities,
        "query": query,
        "gender": gender,
        "university": university,
        "student_id": student_id
    })


@app.get("/admin/students/download/pdf")
async def download_pdf(
    request: Request,
    query: str = "",
    gender: str = "",
    university: str = "",
    student_id: str = ""
):
    await verify_admin(request)

    # Fetch filtered students
    students = await fetch_filtered_students(query, gender, university, student_id)

    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []

    # Title
    styles = getSampleStyleSheet()
    elements.append(Paragraph("Manage Students Report", styles['Title']))

    # Table data
    data = [["ID", "Name", "Email", "Gender", "University"]]
    for student in students:
        gender = student["gender"] if student["gender"] else "Not specified"
        data.append([
            str(student["student_id"]),
            student["name"],
            student["email"],
            gender,
            student["university_name"]
        ])

    # Create table
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(table)

    # Build PDF
    doc.build(elements)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=students_report.pdf"}
    )

# Word Download Route
@app.get("/admin/students/download/word")
async def download_word(
    request: Request,
    query: str = "",
    gender: str = "",
    university: str = "",
    student_id: str = ""
):
    await verify_admin(request)

    # Fetch filtered students
    students = await fetch_filtered_students(query, gender, university, student_id)

    # Create Word document
    doc = Document()
    doc.add_heading("Manage Students Report", 0)

    # Add table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ["ID", "Name", "Email", "Gender", "University"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    for student in students:
        row_cells = table.add_row().cells
        gender = student["gender"] if student["gender"] else "Not specified"
        row_cells[0].text = str(student["student_id"])
        row_cells[1].text = student["name"]
        row_cells[2].text = student["email"]
        row_cells[3].text = gender
        row_cells[4].text = student["university_name"]

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=students_report.docx"}
    )

# CSV Download Route
@app.get("/admin/students/download/csv")
async def download_csv(
    request: Request,
    query: str = "",
    gender: str = "",
    university: str = "",
    student_id: str = ""
):
    await verify_admin(request)

    # Fetch filtered students
    students = await fetch_filtered_students(query, gender, university, student_id)

    # Create CSV
    buffer = io.StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    writer.writerow(["ID", "Name", "Email", "Gender", "University"])

    for student in students:
        gender = student["gender"] if student["gender"] else "Not specified"
        writer.writerow([
            student["student_id"],
            student["name"],
            student["email"],
            gender,
            student["university_name"]
        ])

    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=students_report.csv"}
    )

@app.get("/admin/students/{student_id}/edit", response_class=HTMLResponse)
async def edit_student_form(request: Request, student_id: int):
    async with db_pool.acquire() as conn:
        student = await conn.fetchrow("SELECT * FROM student WHERE student_id = $1", student_id)
        if not student:
            raise HTTPException(status_code=404, detail="Student not found")
    return templates.TemplateResponse("edit_student.html", {"request": request, "student": student})

@app.post("/admin/students/{student_id}/update", response_class=HTMLResponse)
async def update_student(
    request: Request,
    student_id: int,
    name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    gender: str = Form(...)
):
    # Validate gender
    if gender not in ["Male", "Female"]:
        async with db_pool.acquire() as conn:
            student = await conn.fetchrow("SELECT * FROM student WHERE student_id = $1", student_id)
        return templates.TemplateResponse("edit_student.html", {
            "request": request,
            "student": student,
            "error": "Gender must be either Male or Female."
        })

    async with db_pool.acquire() as conn:
        await conn.execute(
            "UPDATE student SET name = $1, email = $2, password = $3, gender = $4 WHERE student_id = $5",
            name, email, password, gender, student_id
        )
    return RedirectResponse(url="/admin/students", status_code=status.HTTP_303_SEE_OTHER)

@app.get("/admin/students/{student_id}/delete", response_class=HTMLResponse)
async def confirm_delete_student(request: Request, student_id: int):
    async with db_pool.acquire() as conn:
        student = await conn.fetchrow("SELECT * FROM student WHERE student_id = $1", student_id)
        if not student:
            raise HTTPException(status_code=404, detail="Student not found")
    return templates.TemplateResponse("confirm_delete.html", {"request": request, "student": student})

@app.post("/admin/students/{student_id}/delete", response_class=HTMLResponse)
async def delete_student(request: Request, student_id: int):
    async with db_pool.acquire() as conn:
        await conn.execute("DELETE FROM registration WHERE student_id = $1", student_id)
        await conn.execute("DELETE FROM student WHERE student_id = $1", student_id)
    return RedirectResponse(url="/admin/students", status_code=status.HTTP_303_SEE_OTHER)

    

@app.get("/admin/registrations", response_class=HTMLResponse)
async def view_registrations(
    request: Request,
    query: str = "",
    gender: str = "",
    university: str = "",
    student_id: str = ""
):
    await verify_admin(request)

    async with db_pool.acquire() as conn:
        # Fetch all universities for the dropdown
        universities = await conn.fetch("SELECT name FROM university ORDER BY name")

    # Fetch filtered registrations
    registrations = await fetch_filtered_registrations(query, gender, university, student_id)

    return templates.TemplateResponse("registration_list.html", {
        "request": request,
        "registrations": registrations,
        "universities": universities,
        "query": query,
        "gender": gender,
        "university": university,
        "student_id": student_id
    })

async def fetch_filtered_registrations(query: str = "", gender: str = "", university: str = "", student_id: str = ""):
    # Build the WHERE clause dynamically
    conditions = []
    params = []
    param_index = 1

    if query:
        conditions.append("(LOWER(s.name) LIKE LOWER($" + str(param_index) + ") OR CAST(s.student_id AS TEXT) LIKE $" + str(param_index) + ")")
        params.append(f"%{query}%")
        param_index += 1

    if gender:
        conditions.append("s.gender = $" + str(param_index))
        params.append(gender)
        param_index += 1

    if university:
        conditions.append("u.name = $" + str(param_index))
        params.append(university)
        param_index += 1

    if student_id:
        conditions.append("s.student_id = $" + str(param_index))
        params.append(int(student_id) if student_id.isdigit() else 0)
        param_index += 1

    where_clause = " AND ".join(conditions)
    if where_clause:
        where_clause = "WHERE " + where_clause

    async with db_pool.acquire() as conn:
        query_sql = f"""
            SELECT r.registration_id, s.student_id AS student_number, s.name AS student_name, s.email, u.name AS home_university,
                   m.name AS module, r.registered_university AS registered_at, r.notes, s.gender
            FROM registration r
            JOIN student s ON r.student_id = s.student_id
            JOIN university u ON s.home_university = u.university_id
            JOIN module m ON r.module_id = m.module_id
            {where_clause}
            ORDER BY r.registration_id
        """
        return await conn.fetch(query_sql, *params)

# Middleware check for admin session
async def verify_admin(request: Request):
    print("Session on /admin:", request.session)  # Debug print
    if not request.session.get("admin_logged_in"):
        raise HTTPException(status_code=403, detail="Not authorized")

@app.get("/edit-registration/{reg_id}", response_class=HTMLResponse)
async def edit_registration(request: Request, reg_id: int):
    async with db_pool.acquire() as conn:
        registration = await conn.fetchrow("""
            SELECT r.registration_id,
                   s.name AS student_name,
                   s.student_id AS student_number,
                   s.email,
                   u.name AS university,
                   m.name AS module,
                   r.notes,
                   s.gender
            FROM registration r
            JOIN student s ON r.student_id = s.student_id
            JOIN university u ON s.home_university = u.university_id
            JOIN module m ON r.module_id = m.module_id
            WHERE r.registration_id = $1
        """, reg_id)
    return templates.TemplateResponse("edit_registration.html", {"request": request, "registration": registration})

@app.post("/edit-registration/{reg_id}")
async def update_registration(
    request: Request,
    reg_id: int,
    full_name: str = Form(...),
    student_id: str = Form(...),
    email: str = Form(...),
    university: str = Form(...),
    module: str = Form(...),
    gender: str = Form(...),
    notes: str = Form("")
):
    gender = gender.title()  # Normalize gender
    if gender not in ["Male", "Female"]:
        async with db_pool.acquire() as conn:
            registration = await conn.fetchrow("""
                SELECT r.registration_id,
                       s.name AS student_name,
                       s.student_id AS student_number,
                       s.email,
                       u.name AS university,
                       m.name AS module,
                       r.notes,
                       s.gender
                FROM registration r
                JOIN student s ON r.student_id = s.student_id
                JOIN university u ON s.home_university = u.university_id
                JOIN module m ON r.module_id = m.module_id
                WHERE r.registration_id = $1
            """, reg_id)
        return templates.TemplateResponse("edit_registration.html", {
            "request": request,
            "registration": registration,
            "error": "Gender must be either Male or Female."
        })

    async with db_pool.acquire() as conn:
        uni = await conn.fetchrow("SELECT university_id FROM university WHERE name = $1", university)
        if not uni:
            uni = await conn.fetchrow("INSERT INTO university (name) VALUES ($1) RETURNING university_id", university)
        university_id = uni["university_id"]

        mod = await conn.fetchrow("SELECT module_id FROM module WHERE name = $1", module)
        if not mod:
            mod = await conn.fetchrow("INSERT INTO module (name, faculty_id) VALUES ($1, 1) RETURNING module_id", module)
        module_id = mod["module_id"]

        await conn.execute(
            "UPDATE student SET name = $1, email = $2, home_university = $3, gender = $4 WHERE student_id = $5",
            full_name, email, university_id, gender, int(student_id)
        )
        await conn.execute(
            "UPDATE registration SET module_id = $1, notes = $2 WHERE registration_id = $3",
            module_id, notes, reg_id
        )

    return RedirectResponse(url="/admin/registrations", status_code=status.HTTP_303_SEE_OTHER)

@app.get("/delete-registration/{reg_id}", response_class=HTMLResponse)
async def confirm_delete_registration(request: Request, reg_id: int):
    async with db_pool.acquire() as conn:
        registration = await conn.fetchrow("""
            SELECT r.registration_id AS id, s.name AS student_name, m.name AS module
            FROM registration r
            JOIN student s ON r.student_id = s.student_id
            JOIN module m ON r.module_id = m.module_id
            WHERE r.registration_id = $1
        """, reg_id)
    return templates.TemplateResponse("confirm_delete.html", {"request": request, "registration": registration})

@app.post("/delete-registration/{reg_id}")
async def delete_registration(reg_id: int):
    async with db_pool.acquire() as conn:
        await conn.execute("DELETE FROM registration WHERE registration_id = $1", reg_id)
    return RedirectResponse(url="/admin/registrations", status_code=status.HTTP_303_SEE_OTHER)

# ------------------- ACADEMIC RESULTS & PROOF -------------------
@app.get("/academic-results", response_class=HTMLResponse)
async def view_results(request: Request):
    student_id = request.session.get("student_id")
    if not student_id:
        return RedirectResponse(url="/tut-login", status_code=303)

    try:
        student_id_int = int(student_id)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="Invalid student ID format")

    async with db_pool.acquire() as conn:
        # Get student details
        student = await conn.fetchrow("""
            SELECT s.student_id, u.name AS home_university
            FROM student s
            JOIN university u ON s.home_university = u.university_id
            WHERE s.student_id = $1
        """, student_id_int)

        if not student:
            raise HTTPException(status_code=404, detail="Student not found")

        # Fetch all registered modules
        registrations = await conn.fetch("""
            SELECT m.module_id, m.name AS module_name, 
                   u.name AS registered_university, u.university_id
            FROM registration r
            JOIN module m ON r.module_id = m.module_id
            JOIN university u ON r.university_id = u.university_id
            WHERE r.student_id = $1
        """, student_id_int)

        results = []

        for reg in registrations:
            # Check if result exists
            result = await conn.fetchrow("""
                SELECT mark, semester, status
                FROM academic_results
                WHERE student_id = $1 AND module_id = $2 AND university_id = $3
            """, student_id_int, reg["module_id"], reg["university_id"])

            if not result:
                # Generate new result if none exists
                mark = random.randint(40, 99)
                semester = random.choice([1, 2])
                status = "Passed" if mark >= 50 else "Failed"

                try:
                    await conn.execute("""
                        INSERT INTO academic_results 
                        (student_id, module_id, university_id, mark, semester, status)
                        VALUES ($1, $2, $3, $4, $5, $6)
                        ON CONFLICT ON CONSTRAINT academic_results_unique 
                        DO UPDATE SET 
                            mark = EXCLUDED.mark,
                            semester = EXCLUDED.semester,
                            status = EXCLUDED.status
                    """, student_id_int, reg["module_id"], reg["university_id"], 
                       mark, semester, status)
                except Exception as e:
                    print("Insert error:", e)
                    raise HTTPException(status_code=500, detail="Error inserting academic result")

                result = {"mark": mark, "semester": semester, "status": status}

            results.append({
                "module_code": f"MOD{reg['module_id']}",
                "module_name": reg["module_name"],
                "your_university": student["home_university"],
                "module_university": reg["registered_university"],
                "mark": f"{result['mark']}%",
                "semester": result["semester"],
                "status": result["status"]
            })

    return templates.TemplateResponse("academic_results.html", {
        "request": request,
        "results": results
    })
    
    
    
    
    
    
@app.get("/academic-results/download/pdf")
async def download_academic_results_pdf(request: Request):
    student_id = request.session.get("student_id")
    if not student_id:
        return RedirectResponse(url="/tut-login", status_code=303)

    try:
        student_id_int = int(student_id)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="Invalid student ID format")

    async with db_pool.acquire() as conn:
        student = await conn.fetchrow("""
            SELECT s.name AS student_name, u.name AS home_university
            FROM student s
            JOIN university u ON s.home_university = u.university_id
            WHERE s.student_id = $1
        """, student_id_int)

        if not student:
            raise HTTPException(status_code=404, detail="Student not found")

        registrations = await conn.fetch("""
            SELECT m.module_id, m.name AS module_name,
                   u.name AS registered_university, u.university_id
            FROM registration r
            JOIN module m ON r.module_id = m.module_id
            JOIN university u ON r.university_id = u.university_id
            WHERE r.student_id = $1
        """, student_id_int)

        results = []
        for reg in registrations:
            result = await conn.fetchrow("""
                SELECT mark, semester, status
                FROM academic_results
                WHERE student_id = $1 AND module_id = $2 AND university_id = $3
            """, student_id_int, reg["module_id"], reg["university_id"])

            if not result:
                continue

            results.append([
                f"MOD{reg['module_id']}",
                reg["module_name"],
                student["home_university"],
                reg["registered_university"],
                f"{result['mark']}%",
                f"Semester {result['semester']}",
                result["status"]
            ])

    # PDF generation
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("EduBridge Academic Results", styles["Title"]))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f"Student Name: {student['student_name']}", styles["Normal"]))
    elements.append(Paragraph(f"Home University: {student['home_university']}", styles["Normal"]))
    elements.append(Spacer(1, 12))

    table_data = [["Module Code", "Module Name", "Home University", "Module University", "Mark", "Semester", "Status"]] + results
    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#4F81BD")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTSIZE', (0,0), (-1,0), 10),
        ('BOTTOMPADDING', (0,0), (-1,0), 12),
        ('BACKGROUND', (0,1), (-1,-1), colors.beige),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
    ]))

    elements.append(table)
    doc.build(elements)
    buffer.seek(0)

    return StreamingResponse(buffer, media_type="application/pdf", headers={
        "Content-Disposition": "attachment; filename=academic_results.pdf"
    })
    
    
    


# ------------------- MODULE INFO AND WEB CRAWLING -------------------
def scrape_universities_for_module_details(module_name, home_university=None):
    module_name = module_name.lower()
    results = []
    
    # University name mapping (common abbreviations to consistent keys)
    UNIVERSITY_ALIASES = {
        'tut': 'TUT',
        'university of pretoria': 'UP',
        'up': 'UP',
        'university of johannesburg': 'UJ',
        'uj': 'UJ',
        'university of witwatersrand': 'WITS',
        'wits': 'WITS',
        'witwatersrand': 'WITS'
    }
    
    # Normalize home university name if provided
    normalized_home_uni = None
    if home_university:
        normalized_home_uni = UNIVERSITY_ALIASES.get(home_university.lower(), home_university).upper()

    # Simulated static responses for specific modules
    simulated_modules = {
        "artificial intelligence": ["WITS", "UP", "UCT", "UKZN", "NWU", "UNISA", "RU", "UWC", "SU"],
        "internet programming": ["TUT", "CPUT", "DUT", "VUT", "CUT", "MUT", "UJ", "UCT", "NWU", "UL", "UNISA", "UKZN"],
        "mobile programming": ["TUT", "CPUT", "DUT", "VUT", "CUT", "MUT", "UJ", "NWU", "UNISA"],
        "software project": ["UCT", "NWU", "UKZN", "UNISA", "TUT", "CPUT", "DUT"],
        "database programming": ["UJ", "UCT", "NWU", "UKZN", "UNISA", "UL", "RU", "TUT", "CPUT", "DUT", "VUT", "CUT", "MUT"],
        "web server management": ["CPUT", "DUT", "VUT", "UJ", "UP", "NWU", "UNISA", "UKZN"],
        "distributed system": ["UP", "WITS", "UCT", "UJ", "NWU", "UKZN", "RU"],
        "data science": ["UP", "WITS", "UCT", "UJ", "NWU", "UKZN", "UNISA", "SU", "RU", "UWC"]
    }
    
    if module_name in simulated_modules:
        for uni in simulated_modules[module_name]:
            # Skip if this is the home university (case-insensitive with normalization)
            if normalized_home_uni and uni.upper() == normalized_home_uni:
                continue
                
            results.append({
                "university": uni,
                "description": "Relevant course found, but no direct snippet.",
                "redirect_url": f"/register?module={module_name.title()}&university={uni}"
            })
        return results
    
    sources = {
        "UJ": "https://www.uj.ac.za/faculties/science ",
        "WITS": "https://www.wits.ac.za/courses/ ",
        "UP": "https://www.up.ac.za/programmes ",
        "TUT": "https://www.tut.ac.za/courses "
    }
    
    for uni, url in sources.items():
        try:
            # Simulate fetching the website
            res = requests.get(url, timeout=5, headers={"User-Agent": "Mozilla/5.0"})
            soup = BeautifulSoup(res.text, "html.parser")
            text = soup.get_text().lower()
            
            # Add keyword variations
            keywords = [module_name]
            if "artificial" in module_name:
                keywords += ["ai", "machine learning", "artificial intelligence and machine learning"]
            
            if any(keyword in text for keyword in keywords):
                index = text.find(keywords[0])
                snippet = text[index:index + 300] if index != -1 else "Relevant course found, but no direct snippet."
                results.append({
                    "university": uni,
                    "description": snippet.strip(),
                    "redirect_url": f"/register?module={module_name.title()}&university={uni}"
                })
        except Exception as e:
            print(f"Error scraping {uni}: {e}")
    
    return results

@app.get("/module-info", response_class=HTMLResponse)
async def show_module_info(request: Request, module: str):
    results = scrape_universities_for_module_details(module)
    return templates.TemplateResponse("module_info.html", {
        "request": request,
        "module": module,
        "results": results
    })

# ------------------- STUDENT REGISTRATION -------------------
def crawl_universities_for_module(module_name: str):
    # Use the updated scrape_universities_for_module_details to get universities
    results = scrape_universities_for_module_details(module_name)
    return [result["university"] for result in results]

@app.get("/register", response_class=HTMLResponse)
async def show_register_form(request: Request, module: str = "", university: str = ""):
    student_id = request.session.get("student_id")
    password = request.session.get("password")
    
    if not student_id or not password:
        return RedirectResponse(url="/university-selection", status_code=302)
    
    # Get student details including gender and home university name
    async with db_pool.acquire() as conn:
        student = await conn.fetchrow("""
            SELECT s.student_id, s.password, s.gender, u.name AS home_university_name
            FROM student s
            JOIN university u ON s.home_university = u.university_id
            WHERE s.student_id = $1
        """, int(student_id))
    
    universities = crawl_universities_for_module(module)
    return templates.TemplateResponse("register.html", {
        "request": request,
        "selected_module": module,
        "universities": universities,
        "selected_university": university,
        "student_id": student["student_id"],
        "password": student["password"],
        "home_university_name": student["home_university_name"],
        "gender": student["gender"]  # Added gender
    })
@app.post("/register")
async def submit_registration(
    request: Request,
    studentID: int = Form(...),
    password: str = Form(...),
    selected_module: str = Form(...),
    chosen_university: str = Form(...),
    gender: str = Form(...),
    comments: str = Form("")
):
    async with db_pool.acquire() as conn:
        # 1. Validate student credentials
        student = await conn.fetchrow("""
            SELECT student_id, name, email, home_university, gender
            FROM student
            WHERE student_id = $1 AND password = $2
        """, studentID, password)

        if not student:
            return templates.TemplateResponse("register.html", {
                "request": request,
                "error": "Invalid student ID or password.",
                "selected_module": selected_module,
                "universities": crawl_universities_for_module(selected_module),
                "student_id": studentID,
                "password": password
            })

        # 2. Check if the student is already registered for this module
        existing_registration = await conn.fetchrow("""
            SELECT 1 FROM registration r
            JOIN module m ON r.module_id = m.module_id
            WHERE r.student_id = $1 AND m.name = $2
        """, studentID, selected_module)

        if existing_registration:
            return templates.TemplateResponse("register.html", {
                "request": request,
                "error": f"You are already registered for {selected_module}.",
                "selected_module": selected_module,
                "universities": crawl_universities_for_module(selected_module),
                "student_id": studentID,
                "password": password,
                "gender": student["gender"]
            })

        # 3. Proceed with registration if not already registered
        home_id = student["home_university"]
        home_university_name = await conn.fetchval(
            "SELECT name FROM university WHERE university_id = $1", home_id
        )

        # 4. Ensure the module exists
        mod = await conn.fetchrow("SELECT module_id FROM module WHERE name = $1", selected_module)
        if not mod:
            mod = await conn.fetchrow(
                "INSERT INTO module (name, faculty_id) VALUES ($1, 1) RETURNING module_id", selected_module
            )
        module_id = mod["module_id"]

        # 5. Get or insert the chosen university
        dest = await conn.fetchrow("SELECT university_id FROM university WHERE name = $1", chosen_university)
        if not dest:
            dest = await conn.fetchrow(
                "INSERT INTO university (name) VALUES ($1) RETURNING university_id", chosen_university
            )
        registered_uni_id = dest["university_id"]

        # 6. Insert registration (will fail if constraint exists)
        try:
            await conn.execute("""
                INSERT INTO registration (student_id, module_id, university_id, notes, registered_university)
                VALUES ($1, $2, $3, $4, $5)
            """, studentID, module_id, registered_uni_id, comments, chosen_university)
        except asyncpg.UniqueViolationError:
            return templates.TemplateResponse("register.html", {
                "request": request,
                "error": f"You are already registered for {selected_module}.",
                "selected_module": selected_module,
                "universities": crawl_universities_for_module(selected_module),
                "student_id": studentID,
                "password": password,
                "gender": student["gender"]
            })

        # 7. Render confirmation
        return templates.TemplateResponse("proof.html", {
            "request": request,
            "student_id": student["student_id"],
            "name": student["name"],
            "email": student["email"],
            "home_university": home_university_name,
            "module": selected_module,
            "registered_university": chosen_university
        })# ------------------- PDF GENERATION -------------------
@app.get("/proof-pdf")
async def generate_proof_pdf(
    request: Request,
    student_id: str,
    name: str,
    email: str,
    home_university: str,
    module: str,
    registered_university: str
):
    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter

    # Header
    c.setFont("Helvetica-Bold", 20)
    if registered_university == "UP":
        c.drawCentredString(width / 2, height - 50, "University of Pretoria")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "Excellence in Education")
    elif registered_university == "UJ":
        c.drawCentredString(width / 2, height - 50, "University of Johannesburg")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "Inspiring Tomorrow")
    elif registered_university == "WITS":
        c.drawCentredString(width / 2, height - 50, "University of the Witwatersrand")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "Knowledge and Insight")
    elif registered_university == "TUT":
        c.drawCentredString(width / 2, height - 50, "Tshwane University of Technology")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "We empower people")
    elif registered_university == "UNIVEN":
        c.drawCentredString(width / 2, height - 50, "University of Venda")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "Creating Future Leaders")
    elif registered_university == "UCT":
        c.drawCentredString(width / 2, height - 50, "University of Cape Town")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "Knowledge for a Better World")
    elif registered_university == "UL":
        c.drawCentredString(width / 2, height - 50, "University of Limpopo")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "Empowering Communities")
    elif registered_university == "NWU":
        c.drawCentredString(width / 2, height - 50, "North-West University")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "It All Starts Here")
    elif registered_university == "RHODES":
        c.drawCentredString(width / 2, height - 50, "Rhodes University")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "Where Leaders Learn")
    elif registered_university == "NMU":
        c.drawCentredString(width / 2, height - 50, "Nelson Mandela University")
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "Change the World")
    else:
        c.drawCentredString(width / 2, height - 50, registered_university)
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, height - 70, "Empowering Future Generations")

    c.line(50, height - 90, width - 50, height - 90)

    # Body
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, height - 130, "TO WHOM IT MAY CONCERN")
    c.setFont("Helvetica", 12)
    c.drawCentredString(width / 2, height - 150, f"It is hereby confirmed that {name}, Student Number {student_id}, is registered from")
    c.drawCentredString(width / 2, height - 170, f"03-Jan-2025 to 14-Dec-2025 at {registered_university} for the under mentioned programme subjects.")

    # Footer
    c.showPage()
    c.save()
    pdf_buffer.seek(0)

    return StreamingResponse(
        pdf_buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment;filename=proof_of_registration_{student_id}.pdf"}
    )

@app.get("/admin/registrations/download/pdf")
async def download_pdf(
    request: Request,
    query: str = "",
    gender: str = "",
    university: str = "",
    student_id: str = ""
):
    await verify_admin(request)

    # Fetch filtered registrations
    registrations = await fetch_filtered_registrations(query, gender, university, student_id)

    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    elements = []

    # Title
    styles = getSampleStyleSheet()
    elements.append(Paragraph("Registered Students Report", styles['Title']))

    # Table data
    data = [["#", "Full Name", "Student ID", "Email", "Home University", "Gender", "Module", "Registered At", "Notes"]]
    for idx, reg in enumerate(registrations, 1):
        gender = reg["gender"] if reg["gender"] else "Not specified"
        data.append([
            str(idx),
            reg["student_name"],
            str(reg["student_number"]),
            reg["email"],
            reg["home_university"],
            gender,
            reg["module"],
            reg["registered_at"],
            reg["notes"]
        ])

    # Create table
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(table)

    # Build PDF
    doc.build(elements)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=registrations_report.pdf"}
    )

# Word Download Route
@app.get("/admin/registrations/download/word")
async def download_word(
    request: Request,
    query: str = "",
    gender: str = "",
    university: str = "",
    student_id: str = ""
):
    try:
        await verify_admin(request)
        
        # Fetch filtered registrations with error handling
        registrations = await fetch_filtered_registrations(query, gender, university, student_id)
        if not registrations:
            raise HTTPException(status_code=404, detail="No registrations found")

        # Create Word document
        doc = Document()
        doc.add_heading('Registered Students Report', 0)

        # Add table with headers
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["#", "Full Name", "Student ID", "Email", "Home University", 
                  "Gender", "Module", "Registered At", "Notes"]
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # Add data rows with null checks
        for idx, reg in enumerate(registrations, 1):
            row_cells = table.add_row().cells
            gender = reg.get("gender", "Not specified") or "Not specified"
            
            row_cells[0].text = str(idx)
            row_cells[1].text = str(reg.get("student_name", ""))
            row_cells[2].text = str(reg.get("student_number", ""))
            row_cells[3].text = str(reg.get("email", ""))
            row_cells[4].text = str(reg.get("home_university", ""))
            row_cells[5].text = str(gender)
            row_cells[6].text = str(reg.get("module", ""))
            row_cells[7].text = str(reg.get("registered_at", ""))
            row_cells[8].text = str(reg.get("notes", ""))

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=registrations_report.docx"}
        )

    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")
# CSV Download Route
@app.get("/admin/registrations/download/csv")
async def download_csv(
    request: Request,
    query: str = "",
    gender: str = "",
    university: str = "",
    student_id: str = ""
):
    await verify_admin(request)

    # Fetch filtered registrations
    registrations = await fetch_filtered_registrations(query, gender, university, student_id)

    # Create CSV
    buffer = io.StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    writer.writerow(["#", "Full Name", "Student ID", "Email", "Home University", "Gender", "Module", "Registered At", "Notes"])

    for idx, reg in enumerate(registrations, 1):
        gender = reg["gender"] if reg["gender"] else "Not specified"
        writer.writerow([
            idx,
            reg["student_name"],
            reg["student_number"],
            reg["email"],
            reg["home_university"],
            gender,
            reg["module"],
            reg["registered_at"],
            reg["notes"]
        ])

    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=registrations_report.csv"}
    )

async def fetch_filtered_students(query: str = "", gender: str = "", university: str = "", student_id: str = ""):
    conditions = []
    params = []
    param_index = 1

    if query:
        conditions.append("(LOWER(s.name) LIKE LOWER($" + str(param_index) + ") OR CAST(s.student_id AS TEXT) LIKE $" + str(param_index) + ")")
        params.append(f"%{query}%")
        param_index += 1

    if gender:
        conditions.append("s.gender = $" + str(param_index))
        params.append(gender)
        param_index += 1

    if university:
        conditions.append("u.name = $" + str(param_index))
        params.append(university)
        param_index += 1

    if student_id:
        conditions.append("s.student_id = $" + str(param_index))
        params.append(int(student_id) if student_id.isdigit() else 0)
        param_index += 1

    where_clause = " AND ".join(conditions)
    if where_clause:
        where_clause = "WHERE " + where_clause

    async with db_pool.acquire() as conn:
        query_sql = f"""
            SELECT s.student_id, s.name, s.email, s.gender, u.name AS university_name
            FROM student s
            JOIN university u ON s.home_university = u.university_id
            {where_clause}
            ORDER BY s.student_id
        """
        return await conn.fetch(query_sql, *params)